import logging

logger = logging.getLogger(__name__)

import fitz
from pathlib import Path
from typing import Optional
from docx import Document

class PDFParserService:
    """Real PDF and DOCX parsing with TXT fallback"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Optional[str]:
        """Extract text from PDF using PyMuPDF with TXT fallback"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            
            if not text.strip():
                raise Exception("PDF is empty")
            
            logger.info(f"✅ PDF parsed: {len(text)} characters from {len(doc)} pages")
            return text.strip()
        
        except Exception as e:
            logger.info(f"⚠️ PDF Parse Error: {e} - trying TXT fallback...")
            # Fallback: read as plain text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read().strip()
                
                if not text:
                    return None
                    
                logger.info(f"✅ TXT fallback SUCCESS: {len(text)} characters")
                return text
            except Exception as txt_err:
                logger.info(f"❌ TXT Fallback FAILED: {txt_err}")
                return None
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Optional[str]:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            result = "\n\n".join(text)
            
            if not result.strip():
                return None
            
            logger.info(f"✅ DOCX parsed: {len(result)} characters, {len(doc.paragraphs)} paragraphs")
            return result.strip()
        
        except Exception as e:
            logger.info(f"❌ DOCX Parse Error: {e}")
            return None
    
    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """Smart extraction based on file extension"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return PDFParserService.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return PDFParserService.extract_text_from_docx(file_path)
        else:
            # Fallback dla innych - czytaj jako TXT
            logger.info(f"⚠️ Unknown extension {ext}, trying TXT...")
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read().strip()
                logger.info(f"✅ TXT read: {len(text)} chars")
                return text if text else None
            except Exception as e:
                logger.info(f"❌ TXT read failed: {e}")
                return None
